from __future__ import print_function
import fitz
import docx2txt
import os
import win32com.client
import glob
from PyPDF2 import PdfFileWriter, PdfFileReader
import textract
import PyPDF2 
from docx import Document
from pyresparser import ResumeParser
import pythoncom
import sys
from docx import *
from docx.shared import Pt



def file_convert(path):
    org_file = []
    if (path.find('.doc') != -1):
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.visible = 0
        pdfs_path = os.getcwd()
        pdfs_path = pdfs_path + "/" # folder where the .pdf files are stored
        for i, doc in enumerate(glob.iglob(pdfs_path+"*.doc")):
            #print(doc)
            filename = doc.split('\\')[-1]
            in_file = os.path.abspath(doc)
            #print(in_file)
            wb = word.Documents.Open(in_file)
            out_file = os.path.abspath(pdfs_path +filename[0:-4]+ ".docx".format(i))
            #print("outfile\n",out_file)
            wb.SaveAs2(out_file, FileFormat=16) # file format for docx
            #print("success...")
            wb.Close()
        word.Quit()
    if(path.find('.pdf') != -1):
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.visible = 0
        pdfs_path = os.getcwd()
        pdfs_path = pdfs_path + "/" # folder where the .pdf files are stored
        for i, doc in enumerate(glob.iglob(pdfs_path+"*.pdf")):
            #print(doc)
            filename = doc.split('\\')[-1]
            in_file = os.path.abspath(doc)
            #print(in_file)
            wb = word.Documents.Open(in_file)
            out_file = os.path.abspath(pdfs_path +filename[0:-4]+ ".docx".format(i))
            #print("outfile\n",out_file)
            wb.SaveAs2(out_file, FileFormat=16) # file format for docx
            #print("success...")
            org_file.append(str(path.replace(".pdf",'')+".pdf"))
            org_file.append(str(path.replace(".pdf",'')+".docx"))
            wb.Close()
        word.Quit()
    elif(path.find('.docx') != -1):
        pythoncom.CoInitialize()
        wdFormatPDF = 17
        in_file = path
        out_file = (str(path.replace(".docx",'')+".pdf"))
        word = win32com.client.DispatchEx("Word.Application")
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        org_file.append(str(path.replace(".docx",'')+".pdf"))
        org_file.append(str(path.replace(".docx",'')+".docx"))
        doc.Close()
        word.Quit()

    return org_file

# name, email, Mobile no.
def person_details(path):
    data = ResumeParser(path).get_extracted_data()
    return data

#word to find linkin id
def linkin(path):
    flag = 0
    temp = docx2txt.process(path)
    text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
    for i in text:
        if(i.find('linkedin')!=-1):
            flag = flag + 1
            return i
    if(flag == 0):
        return "NO Linkin id is found"
    return 0
    

#Number of lines in each page
def no_lines(path):
    count = []
    inputpdf = PdfFileReader(open(path, "rb"))
    for i in range(inputpdf.numPages):
        output = PdfFileWriter()
        output.addPage(inputpdf.getPage(i))
        with open("document.pdf", "wb") as outputStream:
            output.write(outputStream)
        text = str(textract.process("document.pdf"))
        flag =0
        for j in range(len(text)):
            if(text[j] == 'r' and text[j+2] == 'n' and text[j-2] != 'n' and text[j-2] != '7'):
                flag = flag +1
        count.append(flag)
        #print(text)
        os.remove("document.pdf")
    return count

#Number of characters in each page
def no_char(path):
    count = []
    pdfFileObj = open(path, 'rb') 
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 
    for i in range(pdfReader.numPages):
        pageObj = pdfReader.getPage(i) 
        text = (pageObj.extractText())  
        count.append(len(text) - text.count(' '))
    pdfFileObj.close() 
    return count

#find table count
def count_tables(path):
    document = Document(path)
    table_count = 0
    table_count = len(document.tables)
    if(table_count == 0):
        return "Tables Not Found..!"
    return table_count

#number of images present in the Resume
def count_img(path):
    total_img=0
    doc = fitz.open(path)
    for i in range(len(doc)):
        no_img = doc.getPageImageList(i, full=False)
        total_img = total_img + len(no_img)
    if(total_img == 0):
        return "Images Not Found..!"
    doc.close()
    return total_img

def font_size(path):
    document = Document(path)
    size=Pt
    font_size = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if(run.font.size != None):
                size = (run.font.size)
                if size.pt not in font_size:
                    font_size.append(size.pt)
    return font_size
